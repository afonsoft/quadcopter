import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def convert_md_to_docx(md_file_path, docx_file_path):
    """Converte arquivo Markdown para formato DOCX mantendo formatação básica."""
    
    # Ler o arquivo Markdown
    with open(md_file_path, 'r', encoding='utf-8') as file:
        md_content = file.read()
    
    # Criar documento Word
    doc = Document()
    
    # Configurar estilo normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Processar linha por linha
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        # Títulos
        if line.startswith('# '):
            # Título principal
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 0, 0)
                
        elif line.startswith('## '):
            # Título nível 2
            p = doc.add_heading(line[3:], level=2)
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(14)
                
        elif line.startswith('### '):
            # Título nível 3
            p = doc.add_heading(line[4:], level=3)
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(12)
                
        elif line.startswith('#### '):
            # Título nível 4
            p = doc.add_heading(line[5:], level=4)
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                
        elif line.startswith('```'):
            # Início/fim de bloco de código
            continue
            
        elif line.startswith('- '):
            # Lista com marcador
            p = doc.add_paragraph(line[2:], style='List Bullet')
            
        elif line.strip() == '':
            # Linha em branco
            doc.add_paragraph()
            
        elif line.startswith('**') and line.endswith('**'):
            # Texto em negrito
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            
        elif line.startswith('*') and line.endswith('*'):
            # Texto em itálico
            p = doc.add_paragraph()
            run = p.add_run(line[1:-1])
            run.italic = True
            
        elif line.startswith('```python'):
            # Início de bloco de código Python
            continue
            
        elif line.startswith('```'):
            # Fim de bloco de código
            continue
            
        elif line.strip().startswith('```'):
            # Linha de código
            continue
            
        else:
            # Parágrafo normal
            if line.strip():
                # Verificar se é uma linha de código (começa com 4 espaços)
                if line.startswith('    '):
                    p = doc.add_paragraph(line[4:])
                    p.style = doc.styles['Normal']
                    for run in p.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                else:
                    p = doc.add_paragraph(line)
    
    # Salvar documento
    doc.save(docx_file_path)
    print(f"Documento DOCX salvo em: {docx_file_path}")

if __name__ == "__main__":
    md_file = "RELATORIO_COMPLETO.md"
    docx_file = "RELATORIO_COMPLETO.docx"
    
    convert_md_to_docx(md_file, docx_file)
