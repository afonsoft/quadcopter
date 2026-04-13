import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.enum.style import WD_STYLE_TYPE

def convert_md_to_docx_enhanced(md_file_path, docx_file_path):
    """Conversão avançada de Markdown para DOCX com formatação preservada."""
    
    # Ler o arquivo Markdown
    with open(md_file_path, 'r', encoding='utf-8') as file:
        md_content = file.read()
    
    # Criar documento Word
    doc = Document()
    
    # Configurar estilos
    configure_styles(doc)
    
    # Variáveis de controle
    in_code_block = False
    code_block_content = []
    
    # Processar linha por linha
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        # Detectar início/fim de blocos de código
        if line.startswith('```'):
            if in_code_block:
                # Finalizar bloco de código
                if code_block_content:
                    add_code_block(doc, '\n'.join(code_block_content))
                code_block_content = []
                in_code_block = False
            else:
                # Iniciar bloco de código
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line)
            continue
        
        # Processar diferentes tipos de conteúdo
        if process_heading(doc, line):
            continue
        elif process_list(doc, line):
            continue
        elif process_table_row(doc, line):
            continue
        elif process_separator(doc, line):
            continue
        elif process_empty_line(doc, line):
            continue
        else:
            process_paragraph(doc, line)
    
    # Salvar documento
    doc.save(docx_file_path)
    print(f"Documento DOCX avançado salvo em: {docx_file_path}")

def configure_styles(doc):
    """Configura estilos personalizados do documento."""
    
    # Estilo normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Estilo de código
    try:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Consolas'
        code_font.size = Pt(10)
        code_style.paragraph_format.left_indent = Inches(0.25)
    except:
        pass  # Estilo já existe

def process_heading(doc, line):
    """Processa títulos de diferentes níveis."""
    
    if line.startswith('# '):
        # Título principal
        p = doc.add_heading(line[2:], level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_heading_run(p, Pt(16), True)
        return True
        
    elif line.startswith('## '):
        # Título nível 2
        p = doc.add_heading(line[3:], level=2)
        format_heading_run(p, Pt(14), True)
        return True
        
    elif line.startswith('### '):
        # Título nível 3
        p = doc.add_heading(line[4:], level=3)
        format_heading_run(p, Pt(12), True)
        return True
        
    elif line.startswith('#### '):
        # Título nível 4
        p = doc.add_heading(line[5:], level=4)
        format_heading_run(p, Pt(11), True)
        return True
    
    return False

def format_heading_run(paragraph, font_size, bold):
    """Formata os runs de um título."""
    for run in paragraph.runs:
        run.font.bold = bold
        run.font.size = font_size
        run.font.color.rgb = RGBColor(0, 0, 0)

def process_list(doc, line):
    """Processa listas com marcadores."""
    
    if line.startswith('- '):
        p = doc.add_paragraph(line[2:], style='List Bullet')
        return True
    elif line.startswith('* '):
        p = doc.add_paragraph(line[2:], style='List Bullet')
        return True
    elif re.match(r'^\d+\. ', line):
        # Lista numerada
        text = re.sub(r'^\d+\. ', '', line)
        p = doc.add_paragraph(text, style='List Number')
        return True
    
    return False

def process_table_row(doc, line):
    """Processa linhas de tabela (formato Markdown)."""
    
    if '|' in line and not line.startswith('---'):
        # Dividir por pipes e limpar
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        
        if len(cells) >= 2:
            # Adicionar como parágrafo formatado (simplificação)
            p = doc.add_paragraph()
            for i, cell in enumerate(cells):
                run = p.add_run(cell)
                if i < len(cells) - 1:
                    p.add_run(' | ')
            return True
    
    return False

def process_separator(doc, line):
    """Processa separadores horizontais."""
    
    if line.strip() == '---':
        p = doc.add_paragraph()
        p.add_run('_' * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    
    return False

def process_empty_line(doc, line):
    """Processa linhas vazias."""
    
    if line.strip() == '':
        doc.add_paragraph()
        return True
    
    return False

def process_paragraph(doc, line):
    """Processa parágrafos com formatação inline."""
    
    if not line.strip():
        return
    
    # Detectar e processar formatação inline
    p = doc.add_paragraph()
    
    # Usar regex para encontrar diferentes elementos de formatação
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`|[^*`]+)'
    matches = re.findall(pattern, line)
    
    for match in matches:
        if match.startswith('**') and match.endswith('**'):
            # Negrito
            run = p.add_run(match[2:-2])
            run.bold = True
        elif match.startswith('*') and match.endswith('*') and not match.startswith('**'):
            # Itálico
            run = p.add_run(match[1:-1])
            run.italic = True
        elif match.startswith('`') and match.endswith('`'):
            # Código inline
            run = p.add_run(match[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 128)
        else:
            # Texto normal
            p.add_run(match)

def add_code_block(doc, code_text):
    """Adiciona um bloco de código ao documento."""
    
    p = doc.add_paragraph()
    
    # Tentar usar estilo Code, senão usar Normal
    try:
        p.style = doc.styles['Code']
    except:
        p.style = doc.styles['Normal']
    
    # Adicionar linha acima e abaixo do código
    p.add_run('─' * 50 + '\n')
    p.add_run(code_text)
    p.add_run('\n' + '─' * 50)
    
    # Configurar fonte
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)

if __name__ == "__main__":
    md_file = "RELATORIO_COMPLETO.md"
    docx_file = "RELATORIO_COMPLETO_ENHANCED.docx"
    
    convert_md_to_docx_enhanced(md_file, docx_file)
